# src/multimodal/documents.py

import asyncio
from pathlib import Path
from typing import List, Dict, Optional
import PyPDF2
from docx import Document
import pandas as pd
import json
import yaml
import csv
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from odf.opendocument import load
    from odf.text import P

    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False


class DocumentProcessor:
    """
    Обработчик документов для мультимодального мозга
    Читает и анализирует PDF, DOCX, XLSX, CSV, TXT и другие форматы
    """

    def __init__(self, brain, multilingual):
        self.brain = brain
        self.ml = multilingual
        self.download_dir = Path("data/documents")
        self.download_dir.mkdir(parents=True, exist_ok=True)

        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.json': self._process_json,
            '.yaml': self._process_yaml,
            '.yml': self._process_yaml,
            '.md': self._process_markdown
        }

        print("📄 DocumentProcessor инициализирован")
        print(f"   Поддерживается форматов: {len(self.supported_formats)}")

    async def process_document(self, file_path: str) -> Dict:
        """
        Обработка документа

        Args:
            file_path: путь к файлу

        Returns:
            Dict: результаты анализа
        """
        file_path = Path(file_path)
        if not file_path.exists():
            return {'error': 'Файл не найден'}

        ext = file_path.suffix.lower()

        if ext not in self.supported_formats:
            return {
                'error': f'Неподдерживаемый формат: {ext}',
                'supported': list(self.supported_formats.keys())
            }

        # Вызываем соответствующий обработчик
        processor = self.supported_formats[ext]
        results = await processor(file_path)

        # Создаем нейроны на основе документа
        neurons_created = await self._create_document_neurons(results, file_path)
        results['neurons_created'] = neurons_created

        return results

    async def _process_pdf(self, file_path: Path) -> Dict:
        """Обработка PDF файлов"""
        results = {
            'type': 'pdf',
            'pages': 0,
            'text': [],
            'metadata': {},
            'tables': []
        }

        try:
            # Используем pdfplumber если доступен
            if PDFPLUMBER_AVAILABLE:
                with pdfplumber.open(file_path) as pdf:
                    results['pages'] = len(pdf.pages)

                    for i, page in enumerate(pdf.pages[:5]):  # Первые 5 страниц
                        text = page.extract_text()
                        if text:
                            results['text'].append({
                                'page': i + 1,
                                'content': text[:1000]  # Ограничиваем длину
                            })

                        # Извлекаем таблицы
                        tables = page.extract_tables()
                        for table in tables[:2]:  # Первые 2 таблицы
                            if table:
                                results['tables'].append({
                                    'page': i + 1,
                                    'data': table[:5]  # Первые 5 строк
                                })

                    # Метаданные
                    results['metadata'] = {
                        'author': pdf.metadata.get('Author', 'Unknown'),
                        'title': pdf.metadata.get('Title', ''),
                        'subject': pdf.metadata.get('Subject', '')
                    }

            else:
                # Fallback на PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    results['pages'] = len(reader.pages)

                    for i, page in enumerate(reader.pages[:5]):
                        text = page.extract_text()
                        if text:
                            results['text'].append({
                                'page': i + 1,
                                'content': text[:1000]
                            })

        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {e}")

        return results

    async def _process_docx(self, file_path: Path) -> Dict:
        """Обработка DOCX файлов"""
        results = {
            'type': 'docx',
            'paragraphs': 0,
            'text': [],
            'metadata': {}
        }

        try:
            doc = Document(file_path)
            results['paragraphs'] = len(doc.paragraphs)

            # Извлекаем текст
            for i, para in enumerate(doc.paragraphs[:20]):  # Первые 20 параграфов
                if para.text:
                    results['text'].append({
                        'index': i,
                        'content': para.text[:500]
                    })

            # Метаданные
            core_props = doc.core_properties
            results['metadata'] = {
                'author': core_props.author or 'Unknown',
                'title': core_props.title or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }

        except Exception as e:
            logger.error(f"Ошибка обработки DOCX: {e}")

        return results

    async def _process_text(self, file_path: Path) -> Dict:
        """Обработка текстовых файлов"""
        results = {
            'type': 'text',
            'lines': 0,
            'content': '',
            'size': 0
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                results['content'] = content[:2000]  # Первые 2000 символов
                results['lines'] = len(content.split('\n'))
                results['size'] = len(content)

        except UnicodeDecodeError:
            # Пробуем другую кодировку
            with open(file_path, 'r', encoding='cp1251') as f:
                content = f.read()
                results['content'] = content[:2000]
                results['lines'] = len(content.split('\n'))
                results['size'] = len(content)

        return results

    async def _process_csv(self, file_path: Path) -> Dict:
        """Обработка CSV файлов"""
        results = {
            'type': 'csv',
            'rows': 0,
            'columns': 0,
            'headers': [],
            'sample': []
        }

        try:
            df = pd.read_csv(file_path)
            results['rows'] = len(df)
            results['columns'] = len(df.columns)
            results['headers'] = list(df.columns)
            results['sample'] = df.head(5).to_dict('records')

        except Exception as e:
            logger.error(f"Ошибка обработки CSV: {e}")

        return results

    async def _process_excel(self, file_path: Path) -> Dict:
        """Обработка Excel файлов"""
        results = {
            'type': 'excel',
            'sheets': [],
            'total_rows': 0
        }

        try:
            excel_file = pd.ExcelFile(file_path)

            for sheet_name in excel_file.sheet_names[:3]:  # Первые 3 листа
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                results['sheets'].append({
                    'name': sheet_name,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'headers': list(df.columns),
                    'sample': df.head(3).to_dict('records')
                })
                results['total_rows'] += len(df)

        except Exception as e:
            logger.error(f"Ошибка обработки Excel: {e}")

        return results

    async def _process_json(self, file_path: Path) -> Dict:
        """Обработка JSON файлов"""
        results = {
            'type': 'json',
            'structure': {},
            'size': 0
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

                results['structure'] = {
                    'type': type(data).__name__,
                    'keys': list(data.keys()) if isinstance(data, dict) else None,
                    'length': len(data) if isinstance(data, (list, dict)) else None
                }
                results['size'] = len(str(data))
                results['preview'] = str(data)[:500]

        except Exception as e:
            logger.error(f"Ошибка обработки JSON: {e}")

        return results

    async def _process_yaml(self, file_path: Path) -> Dict:
        """Обработка YAML файлов"""
        results = {
            'type': 'yaml',
            'structure': {},
            'size': 0
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)

                results['structure'] = {
                    'type': type(data).__name__,
                    'keys': list(data.keys()) if isinstance(data, dict) else None
                }
                results['size'] = len(str(data))
                results['preview'] = str(data)[:500]

        except Exception as e:
            logger.error(f"Ошибка обработки YAML: {e}")

        return results

    async def _process_markdown(self, file_path: Path) -> Dict:
        """Обработка Markdown файлов"""
        results = {
            'type': 'markdown',
            'headings': [],
            'content': ''
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

                # Извлекаем заголовки
                import re
                headings = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
                results['headings'] = headings
                results['content'] = content[:1000]

        except Exception as e:
            logger.error(f"Ошибка обработки Markdown: {e}")

        return results

    async def _create_document_neurons(self, results: Dict, file_path: Path) -> int:
        """
        Создание нейронов на основе документа
        """
        count = 0
        doc_name = file_path.stem

        # Общий нейрон о документе
        doc_text = f"Документ '{doc_name}' типа {results.get('type', 'unknown')}"

        neurons = self.ml.create_multilingual_neuron(
            self.brain,
            doc_text,
            category=f"Document_{results.get('type', 'file')}"
        )
        count += len(neurons)

        # Нейроны для текстового содержания
        if 'text' in results:
            for text_item in results['text']:
                if isinstance(text_item, dict) and 'content' in text_item:
                    content = text_item['content']
                    if len(content) > 50:
                        neurons = self.ml.create_multilingual_neuron(
                            self.brain,
                            content,
                            category=f"Document_Content"
                        )
                        count += len(neurons)

        # Нейроны для таблиц
        if 'tables' in results:
            for table in results['tables']:
                if 'data' in table and table['data']:
                    table_text = f"Таблица с данными: {str(table['data'][:2])}"
                    neurons = self.ml.create_multilingual_neuron(
                        self.brain,
                        table_text,
                        category=f"Document_Table"
                    )
                    count += len(neurons)

        return count